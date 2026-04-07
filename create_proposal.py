"""대구공공형어린이집연합회 홈페이지 구축 제안서 생성"""
import docx
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ===== 스타일 설정 =====
style = doc.styles['Normal']
style.font.name = 'Pretendard'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.font.color.rgb = RGBColor(0x17, 0x1C, 0x20)

# 한글 폰트 설정
for s in doc.styles:
    try:
        rPr = s.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>')
            rPr.append(rFonts)
        else:
            rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    except:
        pass

PRIMARY = RGBColor(0x00, 0x62, 0x8D)
SECONDARY = RGBColor(0x00, 0x6E, 0x1C)
DARK = RGBColor(0x17, 0x1C, 0x20)
GRAY = RGBColor(0x3F, 0x48, 0x50)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def add_heading_styled(text, level=1, color=PRIMARY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = '맑은 고딕'
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>')
            rPr.append(rFonts)
        else:
            rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return h

def add_para(text, bold=False, color=DARK, size=Pt(11), align=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.font.name = '맑은 고딕'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, header=False):
    row = table.add_row()
    for i, (text, width) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        rPr = run._element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>')
        rPr.append(rFonts)
        if header:
            run.bold = True
            run.font.color.rgb = WHITE
            set_cell_shading(cell, '00628D')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return row

# ===== 페이지 설정 =====
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ===================================================
# 표지
# ===================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('대구공공형어린이집연합회')
run.font.size = Pt(16)
run.font.color.rgb = PRIMARY
run.bold = True
run.font.name = '맑은 고딕'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('공식 홈페이지 구축 제안서')
run.font.size = Pt(28)
run.font.color.rgb = DARK
run.bold = True
run.font.name = '맑은 고딕'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('우리 아이를 위한 최고의 보육 환경,\n신뢰할 수 있는 디지털 플랫폼으로')
run.font.size = Pt(13)
run.font.color.rgb = GRAY
run.font.name = '맑은 고딕'

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026. 04')
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.font.name = '맑은 고딕'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('제안사: (주)제노믹')
run.font.size = Pt(13)
run.font.color.rgb = DARK
run.bold = True
run.font.name = '맑은 고딕'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('대구광역시 서구 국채보상로46길 65-20')
run.font.size = Pt(10)
run.font.color.rgb = GRAY
run.font.name = '맑은 고딕'

doc.add_page_break()

# ===================================================
# 목차
# ===================================================
add_heading_styled('목차', level=1)
doc.add_paragraph()

toc_items = [
    ('1.', '제안 개요', '3'),
    ('2.', '현황 분석 및 필요성', '4'),
    ('3.', '사업 목표 및 기대효과', '5'),
    ('4.', '정보 설계 (IA)', '6'),
    ('5.', '디자인 전략', '7'),
    ('6.', '핵심 기능 상세', '8'),
    ('7.', '기술 아키텍처', '10'),
    ('8.', '추진 일정', '11'),
    ('9.', '투입 인력', '12'),
    ('10.', '사업 비용', '13'),
    ('11.', '유지보수 및 운영 방안', '14'),
    ('12.', '회사 소개', '15'),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {title}')
    run.font.size = Pt(12)
    run.font.name = '맑은 고딕'
    run.font.color.rgb = DARK
    tab = p.add_run('\t' * 6)
    page_run = p.add_run(page)
    page_run.font.size = Pt(12)
    page_run.font.color.rgb = GRAY
    page_run.font.name = '맑은 고딕'

doc.add_page_break()

# ===================================================
# 1. 제안 개요
# ===================================================
add_heading_styled('1. 제안 개요', level=1)

add_heading_styled('1.1 프로젝트 개요', level=2)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

overview_data = [
    ('프로젝트명', '대구공공형어린이집연합회 공식 웹사이트 구축'),
    ('발주처', '대구공공형어린이집연합회'),
    ('제안사', '(주)제노믹'),
    ('사업 기간', '계약일로부터 8주 (약 2개월)'),
    ('주요 대상', '대구시 영유아 학부모, 연합회 소속 원장 및 교사'),
    ('핵심 목표', '대구시 9개 구/군 공공형 어린이집 정보의 투명한 공개 및 접근성 강화'),
]

# 헤더
for i, text in enumerate(['구분', '내용']):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, '00628D')

for label, value in overview_data:
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = '맑은 고딕'
    row.cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

add_heading_styled('1.2 제안 배경', level=2)
add_para('대구광역시에는 약 127개의 공공형 어린이집이 운영되고 있으나, 학부모들이 각 어린이집의 정보를 체계적으로 비교하고 탐색할 수 있는 통합 플랫폼이 부재합니다. 현재 어린이집 정보는 아이사랑포털, 각 어린이집 개별 블로그 등에 분산되어 있어 접근성이 낮고, 연합회 차원의 공지 전달과 회원 간 소통에도 어려움이 있습니다.')
add_para('이에 본 제안서는 대구공공형어린이집연합회의 공식 웹사이트를 구축하여, 학부모에게는 신뢰할 수 있는 정보 탐색 환경을, 연합회 회원에게는 효율적인 소통 채널을 제공하고자 합니다.')

doc.add_page_break()

# ===================================================
# 2. 현황 분석 및 필요성
# ===================================================
add_heading_styled('2. 현황 분석 및 필요성', level=1)

add_heading_styled('2.1 현재 문제점', level=2)

problems = [
    ('정보 분산', '어린이집 정보가 아이사랑포털, 네이버, 카카오맵 등에 흩어져 있어 학부모가 비교 탐색하기 어려움'),
    ('소통 채널 부재', '연합회 공지, 교육 자료, 구인구직 정보를 전달할 공식 채널이 없어 카카오톡 단체방에 의존'),
    ('지역 특화 부족', '전국 단위 포털에서는 대구시 공공형 어린이집만을 필터링하여 보기 어려움'),
    ('모바일 미대응', '기존 연합회 정보 제공 방식이 모바일 환경에 최적화되지 않음'),
]

for title, desc in problems:
    p = doc.add_paragraph()
    run = p.add_run(f'  {title}: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = '맑은 고딕'
    run.font.color.rgb = PRIMARY
    run2 = p.add_run(desc)
    run2.font.size = Pt(11)
    run2.font.name = '맑은 고딕'
    p.paragraph_format.space_after = Pt(8)

add_heading_styled('2.2 기대 효과', level=2)

effects = [
    '학부모의 어린이집 탐색 시간 대폭 단축 (구별 지도 검색)',
    '연합회 공지 전달률 향상 및 소통 효율화',
    '공공형 어린이집의 투명한 정보 공개를 통한 신뢰도 제고',
    '디지털 전환을 통한 연합회 브랜드 이미지 현대화',
    '구인구직 게시판을 통한 보육 인력 수급 원활화',
]

for e in effects:
    p = doc.add_paragraph(e, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = '맑은 고딕'

doc.add_page_break()

# ===================================================
# 3. 사업 목표 및 기대효과
# ===================================================
add_heading_styled('3. 사업 목표', level=1)

goals = [
    ('접근성', '대구시 9개 구/군 어린이집을 인터랙티브 지도에서 직관적으로 탐색할 수 있어야 합니다.'),
    ('신뢰성', '공공형 어린이집의 인증 정보와 최신 현황(정원, 현원, 교직원, 프로그램 등)을 정확하게 전달합니다.'),
    ('편의성', '모바일 기기에서도 모든 기능을 불편 없이 사용할 수 있는 반응형 웹으로 구축합니다.'),
    ('효율성', '관리자가 엑셀 업로드로 어린이집 정보를 일괄 관리할 수 있어야 합니다.'),
]

for i, (title, desc) in enumerate(goals, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {title}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = PRIMARY
    run.font.name = '맑은 고딕'
    p.paragraph_format.space_after = Pt(2)
    add_para(f'    {desc}', size=Pt(11))

doc.add_page_break()

# ===================================================
# 4. 정보 설계 (IA)
# ===================================================
add_heading_styled('4. 정보 설계 (Information Architecture)', level=1)

add_para('사이트의 정보 구조는 사용자의 탐색 목적에 따라 5개 대메뉴로 구성됩니다.')

ia_data = [
    ('홈', '메인 페이지', '히어로 배너, 연합회 현황 통계, 구별 어린이집 지도, 추천 어린이집, 최근 공지사항'),
    ('어린이집 찾기', '검색 및 탐색', '구별 필터, 특성별 필터(영아전담/장애통합/야간연장), 검색, 어린이집 카드 목록, 상세 페이지'),
    ('알림마당', '게시판/커뮤니티', '공지사항, 보육뉴스, 구인구직, 행사일정, 검색, 카테고리 필터'),
    ('연합회 소개', '기관 정보', '회장 인사말, 공공형 어린이집이란, 조직도, 구별 지회, 연혁, 오시는 길'),
    ('마이페이지', '회원 전용', '회원 정보 관리, 즐겨찾기, 알림 설정 (2차 개발)'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
for i, text in enumerate(['메뉴', '구분', '주요 콘텐츠']):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, '00628D')

for menu, cat, content in ia_data:
    row = table.add_row()
    row.cells[0].text = menu
    row.cells[1].text = cat
    row.cells[2].text = content
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'

doc.add_paragraph()
add_heading_styled('4.1 어린이집 상세 페이지 구성', level=2)

detail_items = [
    '히어로 이미지 + 어린이집명 + 인증 배지',
    '현황 카드: 정원/현원, 교직원 수, 통학차량, 연장보육',
    '기본 정보: 주소(복사 기능), 전화번호(바로 전화), 홈페이지 링크',
    '오시는 길: 카카오맵/네이버지도 연동',
    '특별 활동 프로그램 안내',
    '교직원 구성 상세',
    '입소 상담 신청 CTA',
]

for item in detail_items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'

doc.add_page_break()

# ===================================================
# 5. 디자인 전략
# ===================================================
add_heading_styled('5. 디자인 전략: "The Nurturing Atelier"', level=1)

add_para('본 프로젝트의 디자인 철학은 "The Nurturing Atelier(양육의 아틀리에)"입니다. 공공 기관의 구조적 권위와 어린 시절의 따뜻한 유기적 감성을 균형 있게 결합합니다.')

add_heading_styled('5.1 컬러 시스템', level=2)

color_data = [
    ('Primary', '#00628D', 'Deep Sky Blue', '신뢰 — 헤더, 주요 버튼, 기관 인증 요소'),
    ('Secondary', '#006E1C', 'Verdant Growth', '성장/안전 — 성공 상태, 진행 표시'),
    ('Tertiary', '#755700', 'Warm Vitality', '활기 — 강조, 이벤트, 기쁨 요소'),
    ('Surface', '#F6FAFF', 'Soft Sky', '배경 — 차분하고 신뢰감 있는 기본 배경'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
for i, text in enumerate(['역할', '컬러코드', '이름', '용도']):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, '00628D')

for role, code, name, usage in color_data:
    row = table.add_row()
    for j, val in enumerate([role, code, name, usage]):
        row.cells[j].text = val
        for p in row.cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'

doc.add_paragraph()

add_heading_styled('5.2 디자인 원칙', level=2)

principles = [
    ('No-Line Rule', '콘텐츠 영역 구분에 1px 실선 사용 금지. 배경색 차이(Surface 레이어)로 부드럽게 구분하여 현대적이고 접근하기 쉬운 느낌을 줍니다.'),
    ('Glassmorphism Navigation', '상단/하단 네비게이션에 80% 투명도 + 24px 블러를 적용하여, 콘텐츠 이미지가 자연스럽게 비치는 부드러운 UI를 구현합니다.'),
    ('Pill Buttons', '주요 CTA 버튼은 완전 라운드(9999px)로 처리하여 부드럽고 친근한 인상을 줍니다.'),
    ('Frameless Cards', '카드 컴포넌트에 구분선 없이 여백과 배경색으로만 구조를 표현합니다. 최소 라운드 0.5rem, 기본 1rem.'),
    ('Tonal Layering', '전통적 그림자 대신 Surface 계층(Surface → Surface-Low → Surface-Lowest)으로 깊이감을 표현합니다.'),
    ('큰 글자 + 직관적 버튼', '고령 사용자와 모바일 사용자를 고려하여 가독성 높은 타이포그래피와 터치 영역을 확보합니다.'),
]

for title, desc in principles:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    run.font.color.rgb = PRIMARY
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)
    run2.font.name = '맑은 고딕'
    p.paragraph_format.space_after = Pt(6)

add_heading_styled('5.3 타이포그래피', level=2)
add_para('헤드라인: Plus Jakarta Sans (브랜드 모먼트, 공공 기관의 강인함과 신뢰감)')
add_para('본문: Pretendard / Be Vietnam Pro (정보 전달, 넓은 글꼴 폭으로 규정 문서도 읽기 편함)')

doc.add_page_break()

# ===================================================
# 6. 핵심 기능 상세
# ===================================================
add_heading_styled('6. 핵심 기능 상세', level=1)

add_heading_styled('6.1 구별 어린이집 정보 검색 (핵심 기능)', level=2)

features_search = [
    ('인터랙티브 지도', '대구시 9개 구/군을 시각적으로 표현한 지도 UI. 각 구 클릭 시 해당 지역 어린이집 목록으로 이동합니다.'),
    ('상세 필터', '행정구역, 어린이집명, 특성(영아전담, 장애통합, 야간연장, 휴일보육 등)으로 다중 필터 검색이 가능합니다.'),
    ('어린이집 카드', '사진, 이름, 주소, 정원/현원, 교사 수, 통학버스 여부를 한눈에 보여주는 카드 UI입니다.'),
    ('상세 페이지', '기본정보, 지도(카카오/네이버 연동), 프로그램, 교직원 구성, 입소 상담 신청까지 원스톱으로 제공합니다.'),
]

for title, desc in features_search:
    p = doc.add_paragraph()
    run = p.add_run(f'  {title}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = '맑은 고딕'
    p.paragraph_format.space_after = Pt(2)
    add_para(f'    {desc}', size=Pt(10))

add_heading_styled('6.2 게시판 및 커뮤니티', level=2)

board_features = [
    ('공지사항', '연합회 공문 및 주요 행사 안내. 중요 공지 고정 기능.'),
    ('보육뉴스', '최신 보육 정책 및 대구시 보육 관련 소식 공유.'),
    ('구인구직', '소속 어린이집 전용 채용 게시판. 구별 태그 지원.'),
    ('자료실', '회원(원장/교사) 전용 교육 자료 및 서식 다운로드 공간.'),
]

for title, desc in board_features:
    p = doc.add_paragraph()
    run = p.add_run(f'  {title}: ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)
    run2.font.name = '맑은 고딕'
    p.paragraph_format.space_after = Pt(4)

add_heading_styled('6.3 관리자 기능', level=2)

admin_features = [
    '어린이집 정보 일괄 관리: 엑셀 업로드를 통한 대량 정보 업데이트',
    '게시판 관리: 공지사항, 보육뉴스, 구인구직 CRUD',
    '권한 관리: 각 구별 지회장에게 해당 구역 게시판 관리 권한 부여',
    '통계 대시보드: 방문자 수, 인기 어린이집, 게시판 활동 현황',
]

for item in admin_features:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'

doc.add_page_break()

# ===================================================
# 7. 기술 아키텍처
# ===================================================
add_heading_styled('7. 기술 아키텍처', level=1)

add_heading_styled('7.1 기술 스택', level=2)

tech_data = [
    ('프론트엔드', 'React 19 + TypeScript + Vite 8', '최신 웹 표준, 빠른 빌드, 타입 안전성'),
    ('스타일링', 'Tailwind CSS v4', '유틸리티 퍼스트, 빠른 커스터마이징, 작은 번들 크기'),
    ('라우팅', 'React Router v7', 'SPA 기반 페이지 전환, SEO 대응'),
    ('호스팅', 'Cloudflare Pages', '글로벌 CDN, HTTPS 자동, 무료 SSL, 99.99% 가용성'),
    ('백엔드', 'Cloudflare Workers', '서버리스 API, 저지연, 자동 스케일링'),
    ('데이터베이스', 'Cloudflare D1 + Firestore', '관계형(D1) + NoSQL(Firestore) 하이브리드'),
    ('지도 API', '카카오맵 API', '국내 최적화, 정확한 주소 검색'),
    ('아이콘', 'Material Symbols', '구글 공식 아이콘 시스템, 다양한 스타일'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
for i, text in enumerate(['영역', '기술', '선정 사유']):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, '00628D')

for area, tech, reason in tech_data:
    row = table.add_row()
    for j, val in enumerate([area, tech, reason]):
        row.cells[j].text = val
        for p in row.cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'

doc.add_paragraph()

add_heading_styled('7.2 아키텍처 다이어그램', level=2)
add_para('[ 사용자 브라우저 ]')
add_para('        |')
add_para('  [ Cloudflare CDN / Pages ]  ←  정적 파일 (React SPA)')
add_para('        |')
add_para('  [ Cloudflare Workers ]  ←  API 서버 (서버리스)')
add_para('        |')
add_para('  [ D1 Database / Firestore ]  ←  데이터 저장')

add_heading_styled('7.3 보안 및 성능', level=2)
security_items = [
    'HTTPS 전 구간 적용 (Cloudflare 자동 인증서)',
    '회원 개인정보 암호화 저장',
    '이미지 Lazy Loading으로 초기 로딩 속도 최적화',
    '글로벌 CDN으로 300ms 이내 응답 시간 보장',
    'Cloudflare DDoS 방어 기본 탑재',
]
for item in security_items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'

doc.add_page_break()

# ===================================================
# 8. 추진 일정
# ===================================================
add_heading_styled('8. 추진 일정', level=1)

add_para('총 사업 기간: 8주 (약 2개월)')
doc.add_paragraph()

schedule = [
    ('1주', '요구사항 확정', '기능 명세 확정, 어린이집 데이터 수집, IA 확정'),
    ('2주', '디자인', 'UI/UX 디자인, 디자인 시스템 확정, 시안 검토'),
    ('3~4주', '프론트엔드 개발', '메인, 검색, 상세, 게시판, 소개 페이지 개발'),
    ('5~6주', '백엔드 개발', 'API 구축, DB 설계, 관리자 기능, 데이터 연동'),
    ('7주', '통합 테스트', '크로스브라우저 테스트, 모바일 테스트, 성능 최적화'),
    ('8주', '오픈 및 안정화', '운영 서버 배포, 데이터 이관, 교육, 모니터링'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
for i, text in enumerate(['기간', '단계', '주요 활동']):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, '00628D')

for period, phase, tasks in schedule:
    row = table.add_row()
    for j, val in enumerate([period, phase, tasks]):
        row.cells[j].text = val
        for p in row.cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = '맑은 고딕'

doc.add_page_break()

# ===================================================
# 9. 투입 인력
# ===================================================
add_heading_styled('9. 투입 인력', level=1)

staff = [
    ('PM / 기획', '1명', '프로젝트 관리, 요구사항 분석, 일정 관리'),
    ('UI/UX 디자이너', '1명', '디자인 시스템, 화면 설계, 프로토타입'),
    ('프론트엔드 개발', '1명', 'React SPA 개발, 반응형 구현'),
    ('백엔드 개발', '1명', 'API 개발, DB 설계, 관리자 기능'),
    ('QA / 테스트', '1명', '테스트 계획, 실행, 품질 보증'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
for i, text in enumerate(['역할', '인원', '담당 업무']):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, '00628D')

for role, count, task in staff:
    row = table.add_row()
    for j, val in enumerate([role, count, task]):
        row.cells[j].text = val
        for p in row.cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = '맑은 고딕'

doc.add_page_break()

# ===================================================
# 10. 사업 비용
# ===================================================
add_heading_styled('10. 사업 비용', level=1)

add_para('※ 아래 금액은 제안 기준이며, 협의를 통해 조정 가능합니다.', color=GRAY, size=Pt(10))
doc.add_paragraph()

costs = [
    ('UI/UX 디자인', '디자인 시스템, 5개 페이지 시안, 반응형', '3,000,000'),
    ('프론트엔드 개발', 'React SPA, 5개 페이지, 반응형', '5,000,000'),
    ('백엔드 개발', 'API, DB, 관리자 기능, 데이터 연동', '4,000,000'),
    ('지도 API 연동', '카카오맵 연동, 어린이집 위치 표시', '1,000,000'),
    ('테스트 및 QA', '크로스브라우저, 모바일, 성능 테스트', '1,000,000'),
    ('PM / 기획', '프로젝트 관리, 문서화, 교육', '1,000,000'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
for i, text in enumerate(['항목', '상세', '금액 (원, VAT 별도)']):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, '00628D')

for item, detail, amount in costs:
    row = table.add_row()
    for j, val in enumerate([item, detail, amount]):
        row.cells[j].text = val
        for p in row.cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = '맑은 고딕'
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

# 합계
row = table.add_row()
row.cells[0].text = '합계'
row.cells[1].text = ''
row.cells[2].text = '15,000,000'
for cell in row.cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = '맑은 고딕'
    set_cell_shading(cell, 'F0F4F9')
row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()

add_heading_styled('10.1 운영 비용 (월)', level=2)
ops_costs = [
    ('호스팅 (Cloudflare)', '무료 (Pages Free 플랜)'),
    ('도메인', '연 20,000원 (.or.kr 기준)'),
    ('카카오맵 API', '무료 (일 30만 콜 이내)'),
    ('유지보수 (선택)', '월 300,000원 (경미한 수정, 모니터링)'),
]
for item, cost in ops_costs:
    p = doc.add_paragraph()
    run = p.add_run(f'  {item}: ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    run2 = p.add_run(cost)
    run2.font.size = Pt(10)
    run2.font.name = '맑은 고딕'

doc.add_page_break()

# ===================================================
# 11. 유지보수 및 운영 방안
# ===================================================
add_heading_styled('11. 유지보수 및 운영 방안', level=1)

add_heading_styled('11.1 무상 하자보수', level=2)
add_para('납품 후 3개월간 무상 하자보수를 제공합니다. 개발 결함에 의한 오류 수정, 브라우저 호환성 이슈 대응을 포함합니다.')

add_heading_styled('11.2 유지보수 계약 (선택)', level=2)
maintenance = [
    '월 1회 정기 점검 (보안 패치, 성능 모니터링)',
    '경미한 콘텐츠 수정 (월 4건 이내)',
    '긴급 장애 대응 (24시간 이내)',
    '분기 1회 기능 개선 제안 리포트',
]
for item in maintenance:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'

add_heading_styled('11.3 향후 확장 로드맵', level=2)
roadmap = [
    ('2차 개발', '회원 로그인/마이페이지, 관리자 권한 관리, 엑셀 일괄 업로드'),
    ('3차 개발', '입소 대기 현황 실시간 연동, 모바일 앱 푸시 알림'),
    ('4차 개발', '학부모 리뷰/평가 시스템, AI 어린이집 추천'),
]
for phase, desc in roadmap:
    p = doc.add_paragraph()
    run = p.add_run(f'  {phase}: ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    run.font.color.rgb = PRIMARY
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)
    run2.font.name = '맑은 고딕'

doc.add_page_break()

# ===================================================
# 12. 회사 소개
# ===================================================
add_heading_styled('12. 회사 소개', level=1)

add_heading_styled('(주)제노믹', level=2, color=DARK)

company_info = [
    ('상호', '(주)제노믹'),
    ('대표자', '김창훈'),
    ('사업자등록번호', '392-88-01401'),
    ('소재지', '대구광역시 서구 국채보상로46길 65-20'),
    ('설립일', '2024년'),
    ('주요 사업', 'AI 돌봄 서비스, 어린이집 디지털 솔루션, 웹/앱 개발'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
for i, text in enumerate(['구분', '내용']):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, '00628D')

for label, value in company_info:
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = '맑은 고딕'
    row.cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

add_heading_styled('12.1 주요 실적', level=2)

portfolio = [
    'genomic.cc — 어린이집/요양원 통합 디지털 플랫폼 (100+ 시설 운영)',
    'hi.genomic.cc — AI 돌봄 서비스 (독거노인 AI 말벗, Guardian 안전 감시)',
    'ai.genomic.cc — AI 교육 플랫폼 (교사용 AI 도구)',
    'photo.genomic.cc — 어린이집 포토부스 서비스',
    'QR 경품 이벤트 플랫폼 — 피크 15,000+ TPS 대용량 처리',
]
for item in portfolio:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'

doc.add_paragraph()

add_heading_styled('12.2 핵심 역량', level=2)
add_para('어린이집 도메인 전문성: 100개 이상의 어린이집/복지시설 홈페이지 운영 경험으로 보육 현장의 니즈를 깊이 이해하고 있습니다. 블로그 자동 연동, AI 상담, 카카오 알림톡 등 어린이집 특화 솔루션을 다수 보유하고 있습니다.')
add_para('Cloudflare 인프라 전문성: Pages, Workers, D1, R2 등 Cloudflare 생태계를 활용한 서버리스 아키텍처 구축에 풍부한 경험을 보유하고 있어, 안정적이면서도 운영 비용이 매우 낮은 시스템을 제공합니다.')

doc.add_paragraph()
doc.add_paragraph()

# 마무리
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('본 제안서의 내용은 협의를 통해 조정 가능하며,\n프로젝트의 성공적인 수행을 위해 최선을 다하겠습니다.')
run.font.size = Pt(11)
run.font.color.rgb = GRAY
run.font.name = '맑은 고딕'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('감사합니다.')
run.font.size = Pt(14)
run.font.color.rgb = PRIMARY
run.bold = True
run.font.name = '맑은 고딕'

# ===== 저장 =====
output_path = r'C:\Users\admin\daegu-daycare\대구공공형어린이집연합회_홈페이지구축_제안서.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
